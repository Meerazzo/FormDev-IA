from docx import Document

from services.rag.ingestion.parsers.base import BaseParser, ParsedDocument


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        document = Document(file_path)

        paragraphs: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        # Extraction simple des tableaux si le document en contient.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        text = "\n\n".join(paragraphs)

        return ParsedDocument(
            text=text,
            pages=[
                {
                    "page": None,
                    "text": text,
                }
            ],
            metadata={
                "parser": "docx",
                "paragraphs_count": len(paragraphs),
                "tables_count": len(document.tables),
            },
        )
